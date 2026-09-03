"""Reports API — create & retrieve competitive analysis reports."""

from __future__ import annotations

from datetime import datetime
from urllib.parse import quote
from io import BytesIO
from uuid import UUID

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from app.application.dto.report_dto import (
    ReportCreateRequest,
    ReportCreateResponse,
    ReportDetailResponse,
    ReportListResponse,
)
from app.infrastructure.persistence import task_report_runtime

router = APIRouter(prefix="/reports", tags=["reports"])

_reports = task_report_runtime.get_reports()
_tasks = task_report_runtime.get_tasks()


def _persist_reports() -> None:
    task_report_runtime.persist_reports()


def _persist_tasks() -> None:
    task_report_runtime.persist_tasks()


def _content_disposition(filename: str) -> str:
    """Build Content-Disposition that supports non-ASCII filenames."""
    ascii_name = "report.docx"
    encoded = quote(filename)
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{encoded}"


@router.post("", response_model=ReportCreateResponse)
async def create_report(body: ReportCreateRequest) -> ReportCreateResponse:
    from app.interfaces.api.dependencies.workflow import get_workflow_launcher

    launcher = get_workflow_launcher()
    result = await launcher.launch(body)
    return ReportCreateResponse(
        task_id=UUID(result.task_id),
        status=result.status,
        message="分析任务已创建",
    )


@router.get("/{task_id}", response_model=ReportDetailResponse)
async def get_report(task_id: UUID) -> ReportDetailResponse:
    report = _reports.get(str(task_id))
    if not report:
        if str(task_id) in _tasks:
            raise HTTPException(status_code=202, detail="报告正在生成中，请稍后重试")
        raise HTTPException(status_code=404, detail="报告不存在")
    return ReportDetailResponse(**report)


@router.get("/{task_id}/download")
async def download_report(task_id: UUID):
    from docx import Document

    report = _reports.get(str(task_id))
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    markdown = report.get("markdown", "")
    if not markdown:
        raise HTTPException(status_code=404, detail="Word文件不存在或尚未生成")

    evidence_sources = report.get("evidence_sources") or []
    if evidence_sources:
        appendix_lines = ["", "## 附录：证据来源", ""]
        for s in evidence_sources:
            if not isinstance(s, dict) or not s.get("source_id"):
                continue
            sid = s["source_id"]
            title = (s.get("title", "") or "").replace("|", "｜").replace("\n", " ")
            date = s.get("date", "") or "-"
            domain = s.get("domain", "") or s.get("source_type", "") or "-"
            appendix_lines.append(f"- {sid} {title} | {date} | {domain}")
        markdown = markdown + "\n" + "\n".join(appendix_lines)

    try:
        doc = Document()
        for line in markdown.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                continue
            elif line == "---":
                doc.add_paragraph("_" * 40)
            else:
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word文件生成失败: {e}")

    filename = f"竞品分析报告_{str(task_id)[:8]}.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


@router.delete("/{task_id}")
async def delete_report(task_id: UUID) -> dict:
    key = str(task_id)
    found = False
    if key in _reports:
        del _reports[key]
        found = True
    if key in _tasks:
        del _tasks[key]
        found = True
    if not found:
        raise HTTPException(status_code=404, detail="报告不存在")
    _persist_reports()
    _persist_tasks()
    return {"status": "deleted", "task_id": key}


@router.get("", response_model=ReportListResponse)
async def list_reports() -> ReportListResponse:
    reports = list(_reports.values())
    return ReportListResponse(reports=reports, total=len(reports))
